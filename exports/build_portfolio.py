# -*- coding: utf-8 -*-
"""
ShopNest QA Portfolio — generates:
  - output/ShopNest_QA_Portfolio.xlsx (multi-sheet)
  - output/ShopNest_QA_Portfolio.docx (compiled narrative + tables)
  - output/ShopNest_QA_Print.html (open in browser → Print to PDF)

Run: python exports/build_portfolio.py
"""
from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = ROOT / "exports" / "output"
OUT.mkdir(parents=True, exist_ok=True)

# Portfolio owner & dates (interview pack — static)
PORTFOLIO_AUTHOR = "Aparna Patnaik"
PORTFOLIO_AS_OF = date(2025, 5, 10)
PORTFOLIO_EXEC_START = date(2025, 5, 1)
PORTFOLIO_EXEC_END = date(2025, 5, 9)


def _fmt_d(d: date) -> str:
    return d.strftime("%d %B %Y")


def _need_pkg():
    try:
        import openpyxl  # noqa: F401
        from docx import Document  # noqa: F401
    except ImportError:
        print("Installing openpyxl and python-docx…")
        import subprocess

        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "openpyxl", "python-docx", "-q"]
        )


_need_pkg()
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def build_test_cases() -> list[dict]:
    """58+ manual test cases — portfolio / ATS keywords."""
    C: list[dict] = []

    def row(
        tc_id: str,
        module: str,
        title: str,
        pre: str,
        steps: str,
        data: str,
        expected: str,
        priority: str,
        ttype: str,
        req: str,
    ):
        C.append(
            {
                "Test Case ID": tc_id,
                "Module": module,
                "Title": title,
                "Preconditions": pre,
                "Test Steps": steps,
                "Test Data": data,
                "Expected Result": expected,
                "Priority": priority,
                "Test Type": ttype,
                "Requirement ID": req,
            }
        )

    # —— Login ——
    row(
        "TC_LGN_001",
        "Login",
        "Successful login with demo credentials",
        "Browser localStorage can be empty or any; user not logged in.",
        "1) Open login.html | 2) Enter valid email/password | 3) Submit",
        "demo@shopnest.test / Password123",
        "User is redirected to home or next URL; greeting shows name.",
        "High",
        "Positive",
        "FR-AUTH-03",
    )
    row(
        "TC_LGN_002",
        "Login",
        "Empty email validation",
        "On login page.",
        "1) Leave email blank | 2) Enter any password | 3) Submit",
        "email blank",
        "HTML5 / app prevents submit or shows validation (browser dependent).",
        "Medium",
        "Negative",
        "FR-AUTH-03",
    )
    row(
        "TC_LGN_003",
        "Login",
        "Empty password",
        "On login page.",
        "1) Enter valid-format email | 2) Leave password blank | 3) Submit",
        "demo@shopnest.test / password blank",
        "Validation prevents silent success.",
        "Medium",
        "Negative",
        "FR-AUTH-03",
    )
    row(
        "TC_LGN_004",
        "Login",
        "Wrong password",
        "On login page.",
        "1) Enter correct email | 2) Enter wrong password | 3) Submit",
        "demo@shopnest.test / WrongPass999",
        "Clear message: invalid credentials (not misleading lock) — see DEF-001.",
        "High",
        "Negative",
        "FR-AUTH-04",
    )
    row(
        "TC_LGN_005",
        "Login",
        "Unknown email",
        "On login page.",
        "1) Enter unknown email | 2) Enter any password | 3) Submit",
        "nope@example.com / Password123",
        "Appropriate invalid login messaging.",
        "Medium",
        "Negative",
        "FR-AUTH-04",
    )
    row(
        "TC_LGN_006",
        "Login",
        "Email case insensitivity (happy)",
        "Valid account exists from demo.",
        "1) Enter DEMO@SHOPNEST.TEST | 2) Correct password | 3) Submit",
        "DEMO@SHOPNEST.TEST / Password123",
        "Login succeeds (case-insensitive email compare).",
        "Low",
        "Positive",
        "FR-AUTH-03",
    )
    row(
        "TC_LGN_007",
        "Login",
        "SQL-injection style string in password (negative / security awareness)",
        "On login page.",
        "1) Valid email | 2) Password ' OR 1=1 -- | 3) Submit",
        "demo@shopnest.test / ' OR 1=1 --",
        "Login fails; no crash; no misleading success.",
        "Low",
        "Negative",
        "FR-AUTH-03",
    )
    row(
        "TC_LGN_008",
        "Login",
        "Session persists after navigation",
        "Logged in.",
        "1) Login | 2) Open products.html | 3) Refresh once",
        "demo credentials",
        "Still shows logged-in state in header.",
        "Medium",
        "Positive",
        "FR-AUTH-05",
    )

    # —— Registration ——
    row(
        "TC_REG_001",
        "Registration",
        "Register with valid data",
        "Logged out.",
        "1) Open register.html | 2) Fill valid fields | 3) Submit",
        "Alex Tester / alex@example.com / Password99 / Password99",
        "Success message; redirect home; user session created.",
        "High",
        "Positive",
        "FR-AUTH-01",
    )
    row(
        "TC_REG_002",
        "Registration",
        "Invalid email format",
        "On register page.",
        "1) Name OK | 2) bad email | 3) valid passwords | Submit",
        "name / not-an-email / Password99 / Password99",
        "Inline error: valid email required.",
        "High",
        "Negative / EP",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_003",
        "Registration",
        "Password < 8 chars",
        "On register page.",
        "1) Fill valid name/email | 2) password 7 chars | Submit",
        "Short12 / short@x.com / Pass12! / Pass12!",
        "Error: minimum 8 characters.",
        "High",
        "Negative / BVA",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_004",
        "Registration",
        "Password exactly 8 chars (boundary)",
        "On register page.",
        "1) Valid fields with password length 8 | Submit",
        "EightChr / eight@x.com / Passw123 / Passw123",
        "Account created.",
        "Medium",
        "Positive / BVA",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_005",
        "Registration",
        "Password mismatch",
        "On register page.",
        "1) Two different passwords | Submit",
        "x / x@x.com / Password99 / Password98",
        "Error: passwords do not match.",
        "High",
        "Negative",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_006",
        "Registration",
        "Name too short",
        "On register page.",
        "1) Name 1 letter | rest valid | Submit",
        "A / a@a.com / Password99 / Password99",
        "Error on name.",
        "Medium",
        "Negative",
        "FR-AUTH-01",
    )
    row(
        "TC_REG_007",
        "Registration",
        "Whitespace-only name",
        "On register page.",
        "1) Name spaces only | rest valid | Submit",
        "'   ' / b@b.com / Password99 / Password99",
        "Validation error (treated as empty/short).",
        "Low",
        "Negative",
        "FR-AUTH-01",
    )
    row(
        "TC_REG_008",
        "Registration",
        "Unicode in display name",
        "On register page.",
        "1) Name with accented chars | valid email/pass | Submit",
        "José QA / jose@example.com / Password99 / Password99",
        "Accepts or rejects consistently with message (document actual).",
        "Low",
        "Positive",
        "FR-AUTH-01",
    )
    row(
        "TC_REG_009",
        "Registration",
        "Leading/trailing spaces in email trimmed?",
        "On register page.",
        "1) Enter '  alex@example.com  ' | valid pass | Submit",
        "spaces around email",
        "Document trim behaviour; no duplicate ghost accounts server-side N/A.",
        "Low",
        "Negative",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_010",
        "Registration",
        "EP — valid partition sample 2",
        "On register page.",
        "1) Second valid email | Submit",
        "Sam / sam.qa@mail.com / Password99 / Password99",
        "Account created.",
        "Medium",
        "Positive / EP",
        "FR-AUTH-01",
    )
    row(
        "TC_REG_011",
        "Registration",
        "EP — invalid partition: double @",
        "On register page.",
        "1) bad@bad@com | Submit",
        "bad@bad@com",
        "Rejected.",
        "Medium",
        "Negative / EP",
        "FR-AUTH-02",
    )
    row(
        "TC_REG_012",
        "Registration",
        "Very long name (BVA stress)",
        "On register page.",
        "1) 120 char name | valid email/pass",
        "120 x 'N'",
        "Graceful handling (truncate, warn, or accept).",
        "Low",
        "Boundary",
        "FR-AUTH-01",
    )

    # —— Catalog ——
    row(
        "TC_CAT_001",
        "Catalog",
        "All products visible on catalog page",
        "Open products.html without query.",
        "1) Load page | 2) Count cards",
        "n/a",
        "8 products shown (demo dataset).",
        "High",
        "Smoke",
        "FR-CAT-01",
    )
    row(
        "TC_CAT_002",
        "Catalog",
        "Price format consistent",
        "Catalog loaded.",
        "1) Scan prices",
        "n/a",
        "All use $xx.xx pattern.",
        "Medium",
        "Sanity",
        "FR-CAT-01",
    )
    row(
        "TC_CAT_003",
        "Catalog",
        "Add default qty 1",
        "Catalog loaded.",
        "1) Leave qty 1 | Add earbuds",
        "p1 qty 1",
        "Badge +1; line in cart qty 1.",
        "High",
        "Positive",
        "FR-CAT-02",
    )
    row(
        "TC_CAT_004",
        "Catalog",
        "Add qty 5",
        "Catalog loaded.",
        "1) Set qty 5 on an item | Add",
        "any product",
        "Cart shows combined qty +5 on line or correct aggregate.",
        "Medium",
        "Positive",
        "FR-CAT-02",
    )
    row(
        "TC_CAT_005",
        "Catalog",
        "Qty field min bound",
        "Catalog loaded.",
        "1) Try set qty 0 then add",
        "qty 0",
        "Clamped to >=1 or validation message.",
        "Medium",
        "Negative / BVA",
        "FR-CAT-02",
    )
    row(
        "TC_CAT_006",
        "Catalog",
        "Qty field max 99",
        "Catalog loaded.",
        "1) Set 99 | Add",
        "qty 99",
        "Cart respects 99 max (or documents deviation).",
        "Medium",
        "BVA",
        "FR-CAT-02",
    )
    row(
        "TC_CAT_007",
        "Catalog",
        "Non-numeric qty",
        "Catalog loaded.",
        "1) Clear qty field | type letters | Add",
        "qty 'abc'",
        "Handled without crash; sensible default.",
        "Low",
        "Negative",
        "FR-CAT-02",
    )
    row(
        "TC_CAT_008",
        "Catalog",
        "Keyboard accessibility — Tab reaches Add button",
        "Catalog loaded.",
        "1) Tab through card controls",
        "keyboard",
        "Focus order logical.",
        "Low",
        "UI",
        "NFR-ACC-01",
    )

    # —— Search ——
    row(
        "TC_SRCH_001",
        "Search",
        "Basic substring search lowercase",
        "products.html",
        "1) Search 'usb'",
        "q=usb",
        "USB-C Hub appears in results.",
        "High",
        "Positive",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_002",
        "Search",
        "No results message",
        "products.html",
        "1) Search 'zzzznomatch'",
        "q=zzzznomatch",
        "Empty state message visible.",
        "Medium",
        "Negative",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_003",
        "Search",
        "Clear search shows all",
        "After filtered view.",
        "1) Clear query | reload products.html",
        "empty q",
        "All products return.",
        "Medium",
        "Sanity",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_004",
        "Search",
        "Case consistency — lowercase category",
        "products.html",
        "1) Search 'lifestyle'",
        "lifestyle",
        "Lifestyle items listed.",
        "High",
        "Positive",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_005",
        "Search",
        "Case consistency — uppercase same term",
        "products.html",
        "1) Search 'LIFESTYLE'",
        "LIFESTYLE",
        "Same count as TC_SRCH_004 (currently fails — DEF-005).",
        "High",
        "Negative / EP",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_006",
        "Search",
        "Leading/trailing spaces in query",
        "products.html",
        "1) Search '  yoga  '",
        "spaces",
        "Trims and finds Yoga Mat or shows consistent empty.",
        "Low",
        "Boundary",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_007",
        "Search",
        "SKU search",
        "products.html",
        "1) Search 'SN-HUB-007'",
        "SKU",
        "Hub product appears.",
        "Medium",
        "Positive",
        "FR-SRCH-01",
    )
    row(
        "TC_SRCH_008",
        "Search",
        "Header search forwards to catalog",
        "index.html",
        "1) Type text in header search | Submit",
        "bottle",
        "products.html?q=bottle loads with results.",
        "Medium",
        "Smoke",
        "FR-SRCH-01",
    )

    # —— Cart ——
    row(
        "TC_CART_001",
        "Cart",
        "Empty cart message",
        "Cart cleared.",
        "1) Open cart.html",
        "empty",
        "Empty state + CTA to catalog.",
        "High",
        "Smoke",
        "FR-CART-02",
    )
    row(
        "TC_CART_002",
        "Cart",
        "Single line subtotal",
        "One item price 24.50 qty 1.",
        "1) Add bottle | open cart",
        "24.50 x1",
        "Subtotal $24.50.",
        "High",
        "Positive",
        "FR-CART-01",
    )
    row(
        "TC_CART_003",
        "Cart",
        "Quantity update recalculates",
        "Item in cart.",
        "1) Change qty to 3 | blur field",
        "qty 3",
        "Line + subtotal refresh.",
        "High",
        "Positive",
        "FR-CART-01",
    )
    row(
        "TC_CART_004",
        "Cart",
        "Remove line",
        "Two items in cart.",
        "1) Remove one line",
        "n/a",
        "Line disappears; subtotal updates.",
        "High",
        "Positive",
        "FR-CART-02",
    )
    row(
        "TC_CART_005",
        "Cart",
        "Proceed to checkout CTA visible when non-empty",
        "Non-empty cart.",
        "1) Open cart | observe CTA",
        "n/a",
        "Checkout link enabled/visible.",
        "High",
        "Smoke",
        "FR-CART-02",
    )
    row(
        "TC_CART_006",
        "Cart",
        "Cart badge matches item count",
        "Add multiple qty.",
        "1) Compare header badge to sum of qty",
        "mixed",
        "Counts match.",
        "Medium",
        "Sanity",
        "FR-CAT-02",
    )
    for qty, tag in [
        (0, "below min"),
        (1, "min"),
        (9, "pre-bulk"),
        (10, "bulk threshold"),
        (11, "post-bulk"),
        (99, "max"),
        (100, "above max"),
    ]:
        row(
            f"TC_CART_B_{qty}",
            "Cart",
            f"BVA quantity {qty} ({tag})",
            "Product in cart or set via catalog.",
            f"1) Set line qty to {qty} | observe totals",
            f"qty={qty}",
            "Totals match business rules; note DEF-003 at qty>=10.",
            "High" if qty in (10, 11) else "Medium",
            "BVA",
            "FR-CART-01",
        )

    row(
        "TC_CART_014",
        "Cart",
        "Footer alignment vs home (UI)",
        "Compare pages.",
        "1) Open cart | 2) Open home | 3) Compare footer grid",
        "visual",
        "No horizontal shift (see DEF-004).",
        "Low",
        "UI",
        "NFR-USA-01",
    )

    # —— Checkout ——
    row(
        "TC_CHK_001",
        "Checkout",
        "Redirect when logged out",
        "Logged out; cart non-empty.",
        "1) Hit checkout.html directly",
        "guest",
        "Redirect to login with next param.",
        "High",
        "Negative",
        "FR-AUTH-05",
    )
    row(
        "TC_CHK_002",
        "Checkout",
        "Happy path valid address",
        "Logged in; cart non-empty.",
        "1) Fill valid street/zip | Submit",
        "123 Oak; Cityville; 90210",
        "Navigates to payment.html; data in sessionStorage.",
        "High",
        "Positive",
        "FR-CHK-01",
    )
    row(
        "TC_CHK_003",
        "Checkout",
        "Reject too-short street",
        "Logged in; cart non-empty.",
        "1) Street 'ab' | other valid | Submit",
        "ab",
        "Blocked with alert / inline error.",
        "Medium",
        "Negative",
        "FR-CHK-02",
    )
    row(
        "TC_CHK_004",
        "Checkout",
        "Reject short postal",
        "Logged in.",
        "1) Zip '12' | Submit",
        "12",
        "Blocked.",
        "Medium",
        "Negative / BVA",
        "FR-CHK-02",
    )
    row(
        "TC_CHK_005",
        "Checkout",
        "City whitespace-only (negative)",
        "Logged in.",
        "1) City = spaces | other valid | Submit",
        "'   '",
        "Should block — currently fails (DEF-002).",
        "High",
        "Negative",
        "FR-CHK-02",
    )
    row(
        "TC_CHK_006",
        "Checkout",
        "Optional line2 accepted blank",
        "Logged in.",
        "1) Leave line2 empty | valid others | Submit",
        "blank line2",
        "Continues successfully.",
        "Low",
        "Positive",
        "FR-CHK-01",
    )
    row(
        "TC_CHK_007",
        "Checkout",
        "Long street input",
        "Logged in.",
        "1) 200 char street | Submit",
        "200 chars",
        "Graceful handling.",
        "Low",
        "Boundary",
        "FR-CHK-01",
    )
    row(
        "TC_CHK_008",
        "Checkout",
        "Empty cart shows empty checkout",
        "Logged in; cart empty.",
        "1) checkout.html",
        "empty",
        "Message + link to catalog.",
        "Medium",
        "Sanity",
        "FR-CHK-01",
    )

    # —— Payment ——
    row(
        "TC_PAY_001",
        "Payment",
        "Subtotal displayed matches cart",
        "Logged in; fresh cart totals known.",
        "1) Open payment.html | compare subtotal",
        "known total",
        "Displayed subtotal matches last cart subtotal.",
        "High",
        "Smoke",
        "FR-PAY-01",
    )
    row(
        "TC_PAY_002",
        "Payment",
        "Empty name on card blocked",
        "On payment page.",
        "1) Leave name blank | valid PAN | Submit",
        "blank name",
        "HTML5 prevents or shows validation.",
        "Medium",
        "Negative",
        "FR-PAY-02",
    )
    row(
        "TC_PAY_003",
        "Payment",
        "Successful mock payment",
        "Logged in; checkout done.",
        "1) PAN without 0000 | valid fields | Submit",
        "4242424242424242",
        "Redirect to confirmation; cart cleared.",
        "High",
        "Positive",
        "FR-PAY-02",
    )
    row(
        "TC_PAY_004",
        "Payment",
        "Too-short PAN",
        "On payment page.",
        "1) Enter 10 digits | Submit",
        "1231231231",
        "Error invalid card.",
        "Medium",
        "Negative / BVA",
        "FR-PAY-02",
    )
    row(
        "TC_PAY_005",
        "Payment",
        "PAN length boundary 12",
        "On payment page.",
        "1) Enter 12-digit pattern | Submit",
        "12-digit string",
        "Document pass/fail against rule.",
        "Low",
        "BVA",
        "FR-PAY-02",
    )
    row(
        "TC_PAY_006",
        "Payment",
        "Decline path messaging (0000 rule)",
        "On payment page.",
        "1) PAN containing 0000 | Submit",
        "4111111111000011",
        "Clear decline reason (currently vague — DEF-006).",
        "High",
        "Negative",
        "FR-PAY-03",
    )
    row(
        "TC_PAY_007",
        "Payment",
        "Logged-out guard",
        "Clear session.",
        "1) Open payment.html directly",
        "guest",
        "Redirect to login or cart per rules.",
        "Medium",
        "Negative",
        "FR-PAY-02",
    )
    row(
        "TC_PAY_008",
        "Payment",
        "Expiry field presence",
        "Payment page.",
        "1) Verify fields visible",
        "n/a",
        "Expiry + CVV visible.",
        "Low",
        "Sanity",
        "FR-PAY-01",
    )

    # —— Order confirmation ——
    row(
        "TC_ORD_001",
        "Order",
        "Confirmation shows order id",
        "After successful payment.",
        "1) Read header of confirmation",
        "n/a",
        "ORD-… id displayed.",
        "High",
        "Smoke",
        "FR-ORD-01",
    )
    row(
        "TC_ORD_002",
        "Order",
        "Line items match submitted cart",
        "Known cart before pay.",
        "1) Compare SKUs/names/qty",
        "snapshot",
        "Lists match.",
        "High",
        "Positive",
        "FR-ORD-01",
    )
    row(
        "TC_ORD_003",
        "Order",
        "Direct open without order shows friendly message",
        "Clear last order storage (optional).",
        "1) Open order-confirmation.html cold",
        "none",
        "Friendly empty state / message.",
        "Medium",
        "Negative",
        "FR-ORD-01",
    )
    row(
        "TC_ORD_004",
        "Order",
        "Ship-to echoes checkout city",
        "Completed flow.",
        "1) Verify ship block",
        "city from checkout",
        "City visible in summary.",
        "Medium",
        "Positive",
        "FR-ORD-01",
    )
    row(
        "TC_ORD_005",
        "Order",
        "Refresh confirmation still shows order (same session)",
        "After order placed.",
        "1) F5 on confirmation",
        "n/a",
        "Order still rendered from localStorage.",
        "Low",
        "Sanity",
        "FR-ORD-01",
    )

    # —— Cross-cutting ——
    row(
        "TC_XCUT_001",
        "Cross-cutting",
        "Broken image / asset path smoke",
        "Network tab optional.",
        "1) Load each page | check console for 404 on css/js",
        "n/a",
        "No missing core assets.",
        "High",
        "Smoke",
        "NFR-PER-01",
    )
    row(
        "TC_XCUT_002",
        "Cross-cutting",
        "Logout clears auth for checkout",
        "Logged in.",
        "1) Logout | attempt checkout",
        "n/a",
        "Redirect login.",
        "High",
        "Regression",
        "FR-AUTH-05",
    )
    row(
        "TC_XCUT_003",
        "Cross-cutting",
        "Mobile viewport 375px — header wraps",
        "DevTools responsive.",
        "1) Set 375px width | navigate",
        "375px",
        "No horizontal scroll catastrophe.",
        "Low",
        "UI",
        "NFR-USA-01",
    )
    row(
        "TC_XCUT_004",
        "Cross-cutting",
        "Regression — smoke suite subset",
        "Post any single-line CSS change.",
        "1) Run SM-01..SM-06 checklist from doc",
        "checklist",
        "All pass or defects logged.",
        "High",
        "Regression",
        "TP",
    )
    row(
        "TC_XCUT_005",
        "Cross-cutting",
        "Back button from payment to checkout",
        "Mid-flow.",
        "1) Browser Back",
        "browser",
        "No duplicate charge messaging (demo); state sensible.",
        "Low",
        "Exploratory",
        "NFR-USA-01",
    )

    return C


def build_defects() -> list[dict]:
    return [
        {
            "Defect ID": "DEF-001",
            "Summary": "Invalid login always shows misleading account lock message",
            "Severity": "S2 — Major",
            "Priority": "P2 — High",
            "Status": "Open",
            "Environment": "Chrome 136 / Windows 11 / file:// demo-app",
            "Steps to Reproduce": "1) Open login.html\n2) Enter valid email demo@shopnest.test\n3) Enter wrong password\n4) Submit",
            "Expected Result": "Message clearly states invalid email or password; lock message only if truly locked.",
            "Actual Result": "Message claims account temporarily locked for 30 minutes (incorrect for wrong password).",
            "Requirement / Req ID": "FR-AUTH-04",
        },
        {
            "Defect ID": "DEF-002",
            "Summary": "Checkout accepts whitespace-only city",
            "Severity": "S2 — Major",
            "Priority": "P1 — Immediate",
            "Status": "Open",
            "Environment": "Edge 136 / demo static build",
            "Steps to Reproduce": "1) Login\n2) Add item\n3) checkout.html\n4) Street: 123 Main St\n5) City: three spaces\n6) Zip: 90210\n7) Submit",
            "Expected Result": "Validation error: city required.",
            "Actual Result": "Flow continues to payment.",
            "Requirement / Req ID": "FR-CHK-02",
        },
        {
            "Defect ID": "DEF-003",
            "Summary": "Bulk discount line total off by $0.01 at qty ≥ 10",
            "Severity": "S2 — Major",
            "Priority": "P2 — High",
            "Status": "Open",
            "Environment": "Firefox 138",
            "Steps to Reproduce": "1) Add single SKU with unit price 79.99\n2) Set qty 10 in cart\n3) Compare displayed line total to 79.99*10*0.9 rounded banker's/standard 2dp",
            "Expected Result": "Line total matches precise 10% discount rounding policy.",
            "Actual Result": "Displayed line total is 1 cent away from correct value.",
            "Requirement / Req ID": "FR-CART-01",
        },
        {
            "Defect ID": "DEF-004",
            "Summary": "Cart page footer columns slightly misaligned",
            "Severity": "S4 — Trivial",
            "Priority": "P4 — Low",
            "Status": "Open",
            "Environment": "Chrome / 1920x1080",
            "Steps to Reproduce": "1) Open index.html note footer\n2) Open cart.html with items\n3) Compare footer grid alignment to home",
            "Expected Result": "Pixel-consistent alignment across pages.",
            "Actual Result": "Footer inner grid shifted ~6px on cart page.",
            "Requirement / Req ID": "NFR-USA-01",
        },
        {
            "Defect ID": "DEF-005",
            "Summary": "Search returns different counts for uppercase vs lowercase same keyword",
            "Severity": "S3 — Minor",
            "Priority": "P3 — Medium",
            "Status": "Open",
            "Environment": "Any browser",
            "Steps to Reproduce": "1) products.html?q=lifestyle\n2) Note result count\n3) products.html?q=LIFESTYLE\n4) Compare",
            "Expected Result": "Same matching products regardless of case.",
            "Actual Result": "Uppercase query yields fewer or zero matches for same semantic category.",
            "Requirement / Req ID": "FR-SRCH-01",
        },
        {
            "Defect ID": "DEF-006",
            "Summary": "Payment failure shows vague error for certain PAN patterns",
            "Severity": "S3 — Minor",
            "Priority": "P2 — High",
            "Status": "Open",
            "Environment": "Chrome",
            "Steps to Reproduce": "1) Complete checkout\n2) On payment enter PAN containing substring 0000\n3) Submit",
            "Expected Result": "Decline code / friendly actionable reason.",
            "Actual Result": "Banner: Something went wrong.",
            "Requirement / Req ID": "FR-PAY-03",
        },
        {
            "Defect ID": "DEF-007",
            "Summary": "Order confirmation: refresh after manual storage clear shows blank without graceful copy",
            "Severity": "S3 — Minor",
            "Priority": "P3 — Medium",
            "Status": "Open",
            "Environment": "Chrome",
            "Steps to Reproduce": "1) Place order\n2) DevTools Application → clear localStorage\n3) Refresh confirmation",
            "Expected Result": "Friendly message: order not found + CS link.",
            "Actual Result": "Generic 'No order found' — acceptable but lacks guidance (polish).",
            "Requirement / Req ID": "FR-ORD-01",
        },
        {
            "Defect ID": "DEF-008",
            "Summary": "Typo risk: CVV field accepts >3 chars visually confusing",
            "Severity": "S4 — Trivial",
            "Priority": "P4 — Low",
            "Status": "Open",
            "Environment": "All",
            "Steps to Reproduce": "1) payment.html\n2) Type 4 digit CVV",
            "Expected Result": "Amex vs non-Amex rules documented in UI.",
            "Actual Result": " maxlength=4 without helper text.",
            "Requirement / Req ID": "FR-PAY-01",
        },
        {
            "Defect ID": "DEF-009",
            "Summary": "Accessibility: error banners not always linked to fields",
            "Severity": "S3 — Minor",
            "Priority": "P3 — Medium",
            "Status": "Open",
            "Environment": "NVDA (optional)",
            "Steps to Reproduce": "1) Trigger login error\n2) Tab from banner",
            "Expected Result": "aria-live / focus management per WCAG best practice.",
            "Actual Result": "Static div; screen reader may miss urgency.",
            "Requirement / Req ID": "NFR-ACC-01",
        },
        {
            "Defect ID": "DEF-010",
            "Summary": "Performance not measured — doc gap only",
            "Severity": "S4 — Trivial",
            "Priority": "P4 — Low",
            "Status": "Closed / Won't Fix",
            "Environment": "N/A",
            "Steps to Reproduce": "N/A",
            "Expected Result": "NFR perf metrics for demo.",
            "Actual Result": "Intentionally out of scope for static portfolio.",
            "Requirement / Req ID": "NFR-PER-01",
        },
    ]


def build_rtm_rows() -> list[dict]:
    return [
        {
            "Requirement ID": "FR-AUTH-03",
            "Requirement": "Successful login",
            "Test Cases": "TC_LGN_001, TC_LGN_006",
            "Status": "Pass",
        },
        {
            "Requirement ID": "FR-AUTH-04",
            "Requirement": "Accurate invalid login messaging",
            "Test Cases": "TC_LGN_004, TC_LGN_005",
            "Status": "Fail (DEF-001)",
        },
        {
            "Requirement ID": "FR-CHK-02",
            "Requirement": "Reject invalid address",
            "Test Cases": "TC_CHK_005",
            "Status": "Fail (DEF-002)",
        },
        {
            "Requirement ID": "FR-CART-01",
            "Requirement": "Correct cart totals",
            "Test Cases": "TC_CART_B_10, TC_CART_B_11",
            "Status": "Fail (DEF-003)",
        },
        {
            "Requirement ID": "FR-SRCH-01",
            "Requirement": "Consistent search",
            "Test Cases": "TC_SRCH_004, TC_SRCH_005",
            "Status": "Fail (DEF-005)",
        },
        {
            "Requirement ID": "FR-PAY-03",
            "Requirement": "Clear payment failure",
            "Test Cases": "TC_PAY_006",
            "Status": "Fail (DEF-006)",
        },
    ]


def autosize_columns(ws):
    for col in ws.columns:
        max_len = 0
        letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                max_len = max(max_len, len(str(cell.value)) if cell.value else 0)
            except Exception:
                pass
        ws.column_dimensions[letter].width = min(max_len + 2, 60)


def write_xlsx(cases: list[dict], defects: list[dict], rtm: list[dict]) -> Path:
    path = OUT / "ShopNest_QA_Portfolio.xlsx"
    wb = openpyxl.Workbook()
    # Cover
    ws0 = wb.active
    ws0.title = "Cover"
    ws0["A1"] = "ShopNest — Manual QA Portfolio"
    ws0["A1"].font = Font(size=16, bold=True)
    ws0["A2"] = "Author / QA Owner"
    ws0["B2"] = PORTFOLIO_AUTHOR
    ws0["A3"] = "Pack generated (as of)"
    ws0["B3"] = _fmt_d(PORTFOLIO_AS_OF)
    ws0["A4"] = "Test execution window"
    ws0["B4"] = f"{_fmt_d(PORTFOLIO_EXEC_START)} – {_fmt_d(PORTFOLIO_EXEC_END)}"
    ws0["A5"] = "Test case count"
    ws0["B5"] = len(cases)
    ws0["A6"] = "Includes"
    ws0["B6"] = "SRS summary, Test Plan refs, 50+ cases, RTM, Jira-style defects, Smoke/Sanity/Regression"

    # Test Cases
    ws1 = wb.create_sheet("TestCases")
    headers = list(cases[0].keys())
    ws1.append(headers)
    for cell in ws1[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1D4ED8")
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    for c in cases:
        ws1.append([c[h] for h in headers])
    for row in ws1.iter_rows(min_row=2, max_row=ws1.max_row):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    autosize_columns(ws1)

    # Defects
    ws2 = wb.create_sheet("Defects_Jira_Style")
    dh = list(defects[0].keys())
    ws2.append(dh)
    for cell in ws2[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="F97316")
    for d in defects:
        ws2.append([d.get(h, "") for h in dh])
    autosize_columns(ws2)

    # RTM
    ws3 = wb.create_sheet("RTM")
    rh = list(rtm[0].keys())
    ws3.append(rh)
    for cell in ws3[1]:
        cell.font = Font(bold=True)
    for r in rtm:
        ws3.append([r[h] for h in rh])
    autosize_columns(ws3)

    # Smoke / Sanity / Regression short
    ws4 = wb.create_sheet("Smoke_Sanity_Regression")
    ws4.append(["Type", "ID", "Description", "Expected"])
    rows = [
        ("Smoke", "SM-01", "Home loads", "No blank page"),
        ("Smoke", "SM-02", "Nav links", "No 404"),
        ("Smoke", "SM-03", "Add item", "Badge increments"),
        ("Smoke", "SM-04", "Demo login", "Session ok"),
        ("Smoke", "SM-05", "Checkout reachable", "Form visible"),
        ("Smoke", "SM-06", "Pay success path", "Confirmation"),
        ("Sanity", "SN-01", "Search lifestyle lowercase", "Results"),
        ("Sanity", "SN-02", "Search LIFESTYLE uppercase", "Matches lowercase count post-fix"),
        ("Sanity", "SN-03", "Cart qty edit", "Totals refresh"),
        ("Regression", "RG-01", "Full checklist doc 07", "All ticked or defects raised"),
    ]
    for x in rows:
        ws4.append(list(x))
    autosize_columns(ws4)

    # Severity matrix snapshot
    ws5 = wb.create_sheet("Severity_Priority")
    ws5.append(["Severity", "Meaning", "Example"])
    ws5.append(["S1 Critical", "System unusable", "Any card fails (N/A here)"])
    ws5.append(["S2 Major", "Major feature wrong", "DEF-001,002,003"])
    ws5.append(["S3 Minor", "Degraded UX", "DEF-005,006"])
    ws5.append(["S4 Trivial", "Cosmetic", "DEF-004"])
    ws5.append([])
    ws5.append(["Priority", "Meaning"])
    ws5.append(["P1 Immediate", "Stop ship"])
    ws5.append(["P2 High", "Next sprint"])
    ws5.append(["P3 Medium", "Backlog"])
    ws5.append(["P4 Low", "Polish"])
    autosize_columns(ws5)

    wb.save(path)
    return path


def md_to_plain_paragraphs(text: str) -> list[str]:
    lines = text.splitlines()
    paras: list[str] = []
    buf: list[str] = []
    for line in lines:
        if line.strip() == "":
            if buf:
                paras.append("\n".join(buf))
                buf = []
            continue
        if line.startswith("#"):
            if buf:
                paras.append("\n".join(buf))
                buf = []
            paras.append("__HEAD__:" + line.lstrip("# ").strip())
        else:
            buf.append(line)
    if buf:
        paras.append("\n".join(buf))
    return paras


def write_docx(cases: list[dict], defects: list[dict]) -> Path:
    path = OUT / "ShopNest_QA_Portfolio.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading("ShopNest — Manual QA Portfolio (Word Export)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"QA owner: {PORTFOLIO_AUTHOR}")
    doc.add_paragraph(
        "Compiled bundle for interviews: SRS excerpts, Test Plan, Scenarios, Techniques, RTM, Defects, Test Cases summary."
    )
    p_gen = doc.add_paragraph()
    r_gen = p_gen.add_run(
        f"As of: {_fmt_d(PORTFOLIO_AS_OF)} · Execution window: {_fmt_d(PORTFOLIO_EXEC_START)} – {_fmt_d(PORTFOLIO_EXEC_END)}"
    )
    r_gen.italic = True

    for fname in sorted(DOCS.glob("*.md")):
        if fname.name == "KNOWN_ISSUES.md":
            continue
        doc.add_page_break()
        doc.add_heading(fname.stem.replace("_", " "), level=1)
        raw = fname.read_text(encoding="utf-8")
        for block in md_to_plain_paragraphs(raw):
            if block.startswith("__HEAD__:"):
                doc.add_heading(block.replace("__HEAD__:", ""), level=2)
            else:
                # strip markdown tables rough — keep as monospace paragraph
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", block)
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    doc.add_heading("Appendix A — Defects (Jira-style table)", level=1)
    tbl = doc.add_table(rows=1, cols=6)
    hdr = tbl.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Severity"
    hdr[2].text = "Priority"
    hdr[3].text = "Status"
    hdr[4].text = "Summary"
    hdr[5].text = "Req"
    for d in defects:
        row = tbl.add_row().cells
        row[0].text = d.get("Defect ID", "")
        row[1].text = d.get("Severity", "")
        row[2].text = d.get("Priority", "")
        row[3].text = d.get("Status", "")
        row[4].text = d.get("Summary", "")
        row[5].text = d.get("Requirement / Req ID", "")

    doc.add_page_break()
    doc.add_heading("Appendix B — Test Case Index (first 15)", level=1)
    t2 = doc.add_table(rows=1, cols=5)
    h2 = t2.rows[0].cells
    h2[0].text = "TC ID"
    h2[1].text = "Module"
    h2[2].text = "Title"
    h2[3].text = "Type"
    h2[4].text = "Req"
    for c in cases[:15]:
        r = t2.add_row().cells
        r[0].text = c["Test Case ID"]
        r[1].text = c["Module"]
        r[2].text = c["Title"][:120]
        r[3].text = c["Test Type"]
        r[4].text = c["Requirement ID"]

    note = doc.add_paragraph(
        f"Full {len(cases)} test cases live in the Excel sheet `TestCases` for sorting and filters."
    )
    note.paragraph_format.space_before = Pt(12)

    doc.save(path)
    return path


def write_print_html(cases: list[dict], defects: list[dict]) -> Path:
    path = OUT / "ShopNest_QA_Print.html"
    def esc(s):
        return (
            str(s)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    rows_html = ""
    for c in cases:
        rows_html += "<tr>" + "".join(f"<td>{esc(c[k])}</td>" for k in c) + "</tr>"

    defect_rows = ""
    for d in defects:
        defect_rows += "<tr>" + "".join(f"<td>{esc(d[k])}</td>" for k in d) + "</tr>"

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<title>ShopNest QA Portfolio — Print to PDF</title>
<style>
  @page {{ size: A4; margin: 14mm; }}
  body {{ font-family: Segoe UI, Arial, sans-serif; color: #111; font-size: 10pt; }}
  h1 {{ font-size: 16pt; }}
  h2 {{ font-size: 12pt; margin-top: 16pt; page-break-after: avoid; }}
  table {{ border-collapse: collapse; width: 100%; margin-top: 6pt; page-break-inside: auto; }}
  th, td {{ border: 1px solid #ccc; padding: 4px 6px; vertical-align: top; font-size: 8pt; }}
  th {{ background: #1d4ed8; color: #fff; }}
  .meta {{ color: #444; margin-bottom: 12pt; }}
  tr {{ page-break-inside: avoid; page-break-after: auto; }}
</style>
</head>
<body>
  <h1>ShopNest — QA Portfolio (PDF-ready)</h1>
  <p class="meta">QA owner: {esc(PORTFOLIO_AUTHOR)} · Use browser Print → Save as PDF. As of {_fmt_d(PORTFOLIO_AS_OF)} · Execution: {_fmt_d(PORTFOLIO_EXEC_START)} – {_fmt_d(PORTFOLIO_EXEC_END)} · Total test cases: {len(cases)}.</p>
  <h2>1. Defects (Jira-style)</h2>
  <table><thead><tr>{"".join(f"<th>{esc(k)}</th>" for k in defects[0])}</tr></thead>
  <tbody>{defect_rows}</tbody></table>
  <h2>2. Manual test cases</h2>
  <table><thead><tr>{"".join(f"<th>{esc(k)}</th>" for k in cases[0])}</tr></thead>
  <tbody>{rows_html}</tbody></table>
</body>
</html>"""
    path.write_text(html, encoding="utf-8")
    return path


def main():
    cases = build_test_cases()
    defects = build_defects()
    rtm = build_rtm_rows()
    print(f"Test cases: {len(cases)}")
    x = write_xlsx(cases, defects, rtm)
    d = write_docx(cases, defects)
    h = write_print_html(cases, defects)
    (OUT / "meta.json").write_text(
        json.dumps(
            {
                "qa_owner": PORTFOLIO_AUTHOR,
                "as_of": PORTFOLIO_AS_OF.isoformat(),
                "execution_window": {
                    "start": PORTFOLIO_EXEC_START.isoformat(),
                    "end": PORTFOLIO_EXEC_END.isoformat(),
                },
                "test_case_count": len(cases),
                "defect_count": len(defects),
                "outputs": [str(x), str(d), str(h)],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    print("Wrote:", x)
    print("Wrote:", d)
    print("Wrote:", h)


if __name__ == "__main__":
    main()
